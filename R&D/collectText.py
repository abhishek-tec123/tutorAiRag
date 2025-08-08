import os
import json
import logging
import requests
from io import BytesIO, StringIO
from urllib.parse import urlparse
from typing import List, Union
from langchain_core.documents import Document
from langchain_core.document_loaders.base import BaseLoader
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
import PyPDF2
import docx
import pandas as pd
from bs4 import BeautifulSoup

# Logger setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class MixedFileTypeLoader(BaseLoader):
    def __init__(self, file_paths: List[str], verbose: bool = True):
        self.file_paths = file_paths
        self.verbose = verbose
        if not self.verbose:
            logger.setLevel(logging.WARNING)

    def load(self) -> List[Document]:
        docs = []
        futures = []
        with ThreadPoolExecutor() as executor:
            for path in self.file_paths:
                futures.append(executor.submit(self._load_single, path))
            for future in as_completed(futures):
                result = future.result()
                if result:
                    if isinstance(result, list):
                        docs.extend(result)
                    else:
                        docs.append(result)
        return docs

    def _load_single(self, path: str):
        try:
            is_url = path.startswith("http://") or path.startswith("https://")
            ext = os.path.splitext(urlparse(path).path if is_url else path)[1].lower()

            logger.info(f"Loading: {path}")
            loader_fn = {
                ".pdf": self._load_pdf,
                ".docx": self._load_word,
                ".doc": self._load_word,
                ".txt": self._load_txt,
                ".md": self._load_txt,
                ".html": self._load_html,
                ".htm": self._load_html,
                ".csv": self._load_csv,
                ".json": self._load_json
            }.get(ext)

            if loader_fn:
                f = self._get_file_buffer(path) if is_url else path
                result = loader_fn(f)
                if isinstance(result, list):
                    return result
                else:
                    return Document(page_content=result, metadata={"source": path})
            else:
                logger.warning(f"Unsupported extension: {ext}")
        except Exception as e:
            logger.error(f"Error loading {path}: {e}", exc_info=True)
        return None

    def _get_file_buffer(self, url: str) -> Union[BytesIO, StringIO]:
        logger.info(f"Fetching from URL: {url}")
        response = requests.get(url)
        response.raise_for_status()

        content_type = response.headers.get("Content-Type", "")
        is_binary = "application" in content_type or "octet-stream" in content_type
        return BytesIO(response.content) if is_binary else StringIO(response.text)

    def _load_pdf(self, file: Union[str, BytesIO]) -> str:
        text = []
        with (open(file, "rb") if isinstance(file, str) else file) as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)

    def _load_word(self, file: Union[str, BytesIO]) -> str:
        f = open(file, "rb") if isinstance(file, str) else file
        doc = docx.Document(f)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _load_txt(self, file: Union[str, StringIO]) -> str:
        return open(file, encoding="utf-8").read() if isinstance(file, str) else file.read()

    def _load_html(self, file: Union[str, StringIO]) -> str:
        raw = open(file, encoding="utf-8").read() if isinstance(file, str) else file.read()
        soup = BeautifulSoup(raw, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    def _load_csv(self, file: Union[str, StringIO]) -> List[Document]:
        df = pd.read_csv(file)
        return [
            Document(
                page_content="\n".join(f"{col}: {row[col]}" for col in df.columns),
                metadata={"row": idx}
            )
            for idx, row in df.iterrows()
        ]

    def _load_json(self, file: Union[str, StringIO]) -> List[Document]:
        data = json.load(open(file, encoding="utf-8")) if isinstance(file, str) else json.load(file)
        return [Document(page_content=json.dumps(obj), metadata={}) for obj in data] if isinstance(data, list) else [Document(page_content=json.dumps(data), metadata={})]

def load_documents(file_paths: List[str], verbose: bool = True) -> List[Document]:
    loader = MixedFileTypeLoader(file_paths, verbose=verbose)
    return loader.load()

def split_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(docs)

def embed_chunks(chunks: List[Document], embedding_model) -> List[List[float]]:
    texts = [chunk.page_content for chunk in chunks]
    return embedding_model.embed_documents(texts)

def build_embedding_json(chunks: List[Document], embeddings: List[List[float]]) -> List[dict]:
    result = []
    for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        entry = {
            "index": idx,
            "source": chunk.metadata.get("source", None),
            "embedding": emb,
            "embedding_size": len(emb),
            "chunk_text": chunk.page_content,
            "metadata": chunk.metadata
        }
        result.append(entry)
    return result

# def main():
#     file_paths = ["/Users/abhishek/Desktop/vectorSearch_with_monogdb/docs-pdf/Large Language Model Cheat Sheet.pdf"]
#     docs = load_documents(file_paths)
#     chunks = split_documents(docs, chunk_size=1000, chunk_overlap=200)

#     embedding_model = HuggingFaceEmbeddings(
#         model_name="sentence-transformers/all-MiniLM-L12-v2",
#         model_kwargs={'device': 'cpu'},
#         encode_kwargs={'normalize_embeddings': True}
#     )
#     embeddings = embed_chunks(chunks, embedding_model)
#     embedding_json = build_embedding_json(chunks, embeddings)

#     # Print the JSON structure (pretty)
#     print(json.dumps(embedding_json, indent=2))

# if __name__ == "__main__":
#     main()

